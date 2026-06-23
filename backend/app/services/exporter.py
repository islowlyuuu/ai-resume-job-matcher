from io import BytesIO

from docx import Document

from app.schemas.analysis import AnalysisRead


def build_resume_docx(analysis: AnalysisRead) -> bytes:
    document = Document()
    document.add_heading(analysis.optimized_headline, level=1)
    document.add_paragraph(analysis.optimized_summary)

    document.add_heading("经历要点", level=2)
    for bullet in analysis.rewritten_bullets:
        document.add_paragraph(bullet, style="List Bullet")

    document.add_heading("ATS 关键词", level=2)
    document.add_paragraph("、".join(analysis.ats_keywords))

    document.add_heading("Boss 开场白", level=2)
    for opener in _openers(analysis):
        document.add_paragraph(opener, style="List Number")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(analysis: AnalysisRead) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:
        raise RuntimeError("PDF 导出需要安装 reportlab。") from exc

    buffer = BytesIO()
    page = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"

    y = height - 48
    page.setFont(font_name, 16)
    page.drawString(48, y, analysis.optimized_headline[:46])
    y -= 34
    page.setFont(font_name, 10)

    for line in _pdf_lines(analysis):
        if y < 56:
            page.showPage()
            page.setFont(font_name, 10)
            y = height - 48
        page.drawString(48, y, line[:88])
        y -= 18

    page.save()
    return buffer.getvalue()


def _pdf_lines(analysis: AnalysisRead) -> list[str]:
    lines = [
        "推荐摘要",
        analysis.optimized_summary,
        "",
        "经历要点",
    ]
    lines.extend(f"- {item}" for item in analysis.rewritten_bullets)
    lines.extend(["", "ATS 关键词", "、".join(analysis.ats_keywords), "", "Boss 开场白"])
    lines.extend(f"- {item}" for item in _openers(analysis))
    return lines


def _openers(analysis: AnalysisRead) -> list[str]:
    return [line.strip() for line in analysis.cover_letter.splitlines() if line.strip()]
