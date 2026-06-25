from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.core.database import init_db
from app.main import app
from app.services import snapshot_writer


client = TestClient(app)
init_db()


def test_analyze_text_snapshot_and_docx_export(monkeypatch, tmp_path):
    monkeypatch.setenv("OPENAI_API_KEY", "")
    monkeypatch.setattr(snapshot_writer, "SNAPSHOT_DIR", tmp_path)

    response = client.post(
        "/api/analyses/text",
        json={
            "resume_text": "张明\n前端开发\n熟悉 React、TypeScript、数据看板和接口联调。",
            "job_description": "前端开发工程师\n要求 React、TypeScript，有数据看板和接口联调经验。",
            "output_mode": "boss",
            "provider": "local",
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["job_core_skills"]
    assert payload["covered_keywords"]
    assert payload["ai_provider"] == "local"

    providers = client.get("/api/analyses/providers")
    assert providers.status_code == 200
    assert any(item["id"] == "deepseek" for item in providers.json())

    snapshot = client.post(f"/api/analyses/{payload['id']}/snapshot")
    assert snapshot.status_code == 200
    filename = snapshot.json()["filename"]
    assert Path(tmp_path / filename).exists()

    exported = client.get(f"/api/analyses/{payload['id']}/export/docx")
    assert exported.status_code == 200
    assert exported.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )


def test_extract_resume_text_from_docx(tmp_path):
    resume_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("张明")
    document.add_paragraph("前端开发工程师，熟悉 React、TypeScript、数据看板和接口联调。")
    document.save(resume_path)

    with resume_path.open("rb") as resume_file:
        response = client.post(
            "/api/analyses/extract-resume",
            files={
                "resume": (
                    "resume.docx",
                    resume_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    assert "React" in response.json()["resume_text"]
